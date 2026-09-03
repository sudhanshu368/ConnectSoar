import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_doc():
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("ConnectSoar — Authentication API Documentation")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(16, 88, 160)
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Production-Grade Authentication, Authorization & Session Management")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    subtitle_p.paragraph_format.space_after = Pt(16)

    # Divider line
    doc.add_paragraph("―" * 55)

    # Section 1: Overview & Architecture
    h1 = doc.add_heading("1. Architecture & Security Principles", level=1)
    h1.runs[0].font.color.rgb = RGBColor(16, 88, 160)
    
    p = doc.add_paragraph(
        "ConnectSoar utilizes Supabase Auth for identity management, password hashing, and token issuance, "
        "combined with a zero-trust Spring Boot backend layer responsible for status enforcement, role authorization, "
        "rate limiting, and first-time password reset enforcement."
    )
    
    bullets = [
        "Base URL: http://localhost:8080/api/v1/auth",
        "Clean Token Contract: Responses contain only access_token and refresh_token (no token_type or expires_in).",
        "Zero-Trust Role Enforcement: Frontend roles are never trusted. Initial employee creation forces role=employee, status=active, reset_password=true.",
        "First-Time Login Gatekeeping: If reset_password=true, login returns HTTP 403 (PASSWORD_CHANGE_REQUIRED) with a single-purpose reset token and NO access/refresh tokens.",
        "Account Status Validation: Users with status inactive or suspended are immediately blocked even with valid unexpired JWTs.",
        "Anti-Enumeration Protection: Forgot password and login endpoints never reveal whether an email exists."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph()

    # Helper for adding API sections
    def add_api(name, method, endpoint, description, req_headers, req_body, responses):
        h2 = doc.add_heading(f"{name}", level=2)
        h2.runs[0].font.color.rgb = RGBColor(30, 60, 120)
        
        meta_p = doc.add_paragraph()
        r_m = meta_p.add_run("HTTP Method: ")
        r_m.bold = True
        r_m_val = meta_p.add_run(f"{method}  |  ")
        r_m_val.font.color.rgb = RGBColor(0, 128, 0) if method in ["GET", "POST"] else RGBColor(200, 100, 0)
        r_m_val.bold = True
        
        r_ep = meta_p.add_run("Endpoint: ")
        r_ep.bold = True
        r_ep_val = meta_p.add_run(f"{endpoint}\n")
        r_ep_val.font.color.rgb = RGBColor(16, 88, 160)
        r_ep_val.bold = True

        r_desc = meta_p.add_run(f"Description: {description}")

        # Headers
        if req_headers:
            h_p = doc.add_paragraph()
            h_p.add_run("Request Headers:\n").bold = True
            for k, v in req_headers.items():
                h_p.add_run(f"  • {k}: {v}\n")

        # Request Body
        if req_body:
            doc.add_paragraph("Request Body (JSON):", style='Normal').runs[0].bold = True
            code_p = doc.add_paragraph()
            code_run = code_p.add_run(req_body)
            code_run.font.name = 'Consolas'
            code_run.font.size = Pt(9.5)
            code_p.paragraph_format.left_indent = Inches(0.2)

        # Responses
        doc.add_paragraph("Responses:", style='Normal').runs[0].bold = True
        for title, status_code, body in responses:
            resp_hdr = doc.add_paragraph()
            resp_hdr.paragraph_format.space_before = Pt(4)
            resp_hdr.paragraph_format.space_after = Pt(2)
            
            r_title = resp_hdr.add_run(f"▸ {title} (Status: {status_code})\n")
            r_title.bold = True
            r_title.font.color.rgb = RGBColor(0, 100, 0) if status_code < 400 else RGBColor(180, 40, 40)
            
            resp_code = doc.add_paragraph()
            code_r = resp_code.add_run(body)
            code_r.font.name = 'Consolas'
            code_r.font.size = Pt(9)
            resp_code.paragraph_format.left_indent = Inches(0.2)
            resp_code.paragraph_format.space_after = Pt(8)

        doc.add_paragraph("―" * 55)

    # 1. Login API
    add_api(
        name="1. User Login",
        method="POST",
        endpoint="/api/v1/auth/login",
        description="Authenticates employee or admin. Enforces active status and reset_password verification.",
        req_headers={"Content-Type": "application/json"},
        req_body='{\n  "email": "employee@example.com",\n  "password": "UserPassword123!"\n}',
        responses=[
            ("Case A: Normal Login (reset_password = false)", 200, 
'''{
  "success": true,
  "message": "Login successful.",
  "data": {
    "access_token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIs...",
    "refresh_token": "0b4c8038-d635-4309-a7b6-c677610368a5",
    "user": {
      "id": "c8d62635-4309-450f-a7b6-c677610368a5",
      "email": "employee@example.com",
      "name": "Rahul Kumar",
      "role": "employee",
      "status": "active",
      "department": "Engineering",
      "designation": "Flutter Developer",
      "phone": "+919876543210",
      "image_url": null,
      "reset_password": false,
      "created_at": "2026-08-31T13:30:00Z",
      "updated_at": "2026-08-31T13:30:00Z"
    }
  }
}'''),
            ("Case B: First-Time Login (reset_password = true) — NO ACCESS/REFRESH TOKENS ISSUED", 403,
'''{
  "success": false,
  "error": {
    "code": "PASSWORD_CHANGE_REQUIRED",
    "message": "Password change is required before accessing the application."
  },
  "data": {
    "reset_password": true,
    "password_reset_token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",
    "user": {
      "id": "c8d62635-4309-450f-a7b6-c677610368a5",
      "email": "new.employee@example.com",
      "name": "Jane New",
      "role": "employee"
    }
  }
}'''),
            ("Case C: Invalid Password / Non-existent User", 401,
'''{
  "success": false,
  "error": {
    "code": "INVALID_CREDENTIALS",
    "message": "Invalid email or password."
  }
}'''),
            ("Case D: Inactive / Suspended Account", 403,
'''{
  "success": false,
  "error": {
    "code": "USER_INACTIVE",
    "message": "User account is inactive."
  }
}'''),
            ("Case E: Rate Limit Exceeded (Too many attempts)", 429,
'''{
  "success": false,
  "error": {
    "code": "RATE_LIMITED",
    "message": "Too many requests. Please try again later."
  }
}''')
        ]
    )

    # 2. Refresh Token API
    add_api(
        name="2. Refresh Access Token",
        method="POST",
        endpoint="/api/v1/auth/refresh",
        description="Refreshes access token using a valid refresh token. Rotates tokens securely.",
        req_headers={"Content-Type": "application/json"},
        req_body='{\n  "refresh_token": "0b4c8038-d635-4309-a7b6-c677610368a5"\n}',
        responses=[
            ("Success Response", 200,
'''{
  "success": true,
  "message": "Token refreshed successfully.",
  "data": {
    "access_token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.NEW_ACCESS_TOKEN...",
    "refresh_token": "NEW_ROTATED_REFRESH_TOKEN..."
  }
}'''),
            ("Invalid / Expired Refresh Token", 401,
'''{
  "success": false,
  "error": {
    "code": "TOKEN_INVALID",
    "message": "Invalid or expired refresh token."
  }
}''')
        ]
    )

    # 3. Get Current User API
    add_api(
        name="3. Get Current User Profile (/me)",
        method="GET",
        endpoint="/api/v1/auth/me",
        description="Retrieves live profile of the currently authenticated user.",
        req_headers={
            "Authorization": "Bearer <access_token>",
            "Content-Type": "application/json"
        },
        req_body=None,
        responses=[
            ("Success Response", 200,
'''{
  "success": true,
  "data": {
    "id": "c8d62635-4309-450f-a7b6-c677610368a5",
    "email": "employee@example.com",
    "name": "Rahul Kumar",
    "role": "employee",
    "status": "active",
    "department": "Engineering",
    "designation": "Flutter Developer",
    "phone": "+919876543210",
    "image_url": null,
    "reset_password": false,
    "created_at": "2026-08-31T13:30:00Z",
    "updated_at": "2026-08-31T13:30:00Z"
  }
}'''),
            ("Missing or Invalid Token", 401,
'''{
  "success": false,
  "error": {
    "code": "UNAUTHORIZED",
    "message": "Missing or invalid Authorization header. Expected 'Authorization: Bearer <token>'"
  }
}''')
        ]
    )

    # 4. Change Password API
    add_api(
        name="4. Change Password",
        method="POST",
        endpoint="/api/v1/auth/change-password",
        description="Updates user password. Supports both first-time password setup (using reset_token) and normal password update (using Bearer session). Automatically sets reset_password = false.",
        req_headers={"Content-Type": "application/json"},
        req_body='''// Option A: First-time setup with reset_token
{
  "new_password": "NewSecurePassword456!",
  "confirm_password": "NewSecurePassword456!",
  "reset_token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9..."
}

// Option B: Authenticated session
{
  "current_password": "OldPassword123!",
  "new_password": "NewSecurePassword456!",
  "confirm_password": "NewSecurePassword456!"
}''',
        responses=[
            ("Success Response", 200,
'''{
  "success": true,
  "message": "Password changed successfully.",
  "data": {
    "reset_password": false
  }
}'''),
            ("Passwords Mismatch / Policy Validation Error", 400,
'''{
  "success": false,
  "error": {
    "code": "VALIDATION_ERROR",
    "message": "New password and confirm password do not match."
  }
}''')
        ]
    )

    # 5. Forgot Password API
    add_api(
        name="5. Forgot Password",
        method="POST",
        endpoint="/api/v1/auth/forgot-password",
        description="Initiates secure Supabase password recovery. Always returns a generic message to prevent email enumeration.",
        req_headers={"Content-Type": "application/json"},
        req_body='{\n  "email": "employee@example.com"\n}',
        responses=[
            ("Generic Success Response (Prevents Account Enumeration)", 200,
'''{
  "success": true,
  "message": "If an account exists with this email, password reset instructions have been sent."
}'''),
            ("Rate Limited Response", 429,
'''{
  "success": false,
  "error": {
    "code": "RATE_LIMITED",
    "message": "Too many requests. Please try again later."
  }
}''')
        ]
    )

    # 6. Logout API
    add_api(
        name="6. User Logout",
        method="POST",
        endpoint="/api/v1/auth/logout",
        description="Revokes current application session.",
        req_headers={
            "Authorization": "Bearer <access_token>",
            "Content-Type": "application/json"
        },
        req_body=None,
        responses=[
            ("Success Response", 200,
'''{
  "success": true,
  "message": "Logged out successfully."
}''')
        ]
    )

    # Error Codes Table
    h2 = doc.add_heading("2. Standardized Error Codes Reference", level=2)
    h2.runs[0].font.color.rgb = RGBColor(16, 88, 160)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Error Code"
    hdr_cells[1].text = "HTTP Status"
    hdr_cells[2].text = "Description / When Triggered"

    for cell in hdr_cells:
        set_cell_background(cell, "1058A0")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    errors = [
        ("INVALID_CREDENTIALS", "401 Unauthorized", "Incorrect email or password entered during login."),
        ("UNAUTHORIZED", "401 Unauthorized", "Missing or malformed Authorization Bearer header."),
        ("TOKEN_EXPIRED", "401 Unauthorized", "Access token has expired. Client must refresh token."),
        ("TOKEN_INVALID", "401 Unauthorized", "Access or refresh token is invalid or cryptographically corrupted."),
        ("PASSWORD_CHANGE_REQUIRED", "403 Forbidden", "Employee reset_password is true. Normal tokens blocked until password is changed."),
        ("USER_INACTIVE", "403 Forbidden", "User account is disabled/inactive. All requests rejected."),
        ("USER_SUSPENDED", "403 Forbidden", "User account has been suspended by administration."),
        ("FORBIDDEN", "403 Forbidden", "Insufficient permissions (e.g. employee attempting admin endpoint)."),
        ("RATE_LIMITED", "429 Too Many Requests", "Exceeded request threshold (e.g. >10 login attempts/min)."),
        ("VALIDATION_ERROR", "400 Bad Request", "Request body failed input validation rules.")
    ]

    for code, status, desc in errors:
        row_cells = table.add_row().cells
        row_cells[0].text = code
        row_cells[1].text = status
        row_cells[2].text = desc
        for cell in row_cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    doc.save("ConnectSoar_Authentication_API_Documentation.docx")
    print("Word document updated successfully: ConnectSoar_Authentication_API_Documentation.docx")

if __name__ == "__main__":
    create_doc()
